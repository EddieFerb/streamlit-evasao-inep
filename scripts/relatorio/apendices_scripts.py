 # Script: apendices_scripts.py
 # Gera um documento Word com todos os códigos-fonte dos scripts principais do projeto, com enumeração de linhas.

import os
from docx import Document
from docx.shared import Pt, Inches
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# Lista de caminhos relativos dos scripts
script_paths = [
    "scripts/coleta_dados/coletar_links_inep.py",
    "scripts/coleta_dados/coleta_dados_oficiais.py",
    "scripts/processamento_dados/pre_processamento.py",
    "scripts/processamento_dados/pre_processamento_Transfer_Learn.py",
    "scripts/processamento_dados/tratar_dados.py",
    "scripts/processamento_dados/tratar_dados_Transfer_Learn.py",
    "scripts/processamento_dados/preparar_entrada_modelos.py",
    "scripts/modelagem/gerar_base_modelo.py",
    "scripts/modelagem/modelagem.py",
    "scripts/modelagem/treinamento_modelo_C4.5_Tree_J48.py",
    "scripts/modelagem/feature-based.py.py",
    "scripts/modelagem/treinamento_modelo_Fine-tuning.py",
    "scripts/analises/analises.py",
    "scripts/analises/comparar_modelos.py",
    "scripts/analises/validar_modelos_temporais.py",
    "scripts/analises/prever_com_modelos.py",
    "scripts/analises/resumir_modelo_h5.py",
    "scripts/visualizacao/gerar_graficos.py",
    "scripts/visualizacao/comparar_predicoes_cursos.py",
    "scripts/dashboard/app_evasao.py",
]

# Criação do documento Word
doc = Document()
doc.add_heading("Apêndice - Códigos Python Utilizados", level=1)

for path in script_paths:
    if os.path.exists(path):
        doc.add_page_break()
        doc.add_heading(os.path.basename(path), level=2)
        with open(path, 'r', encoding='utf-8') as f:
            lines = f.readlines()
        for i, line in enumerate(lines, start=1):
            paragraph = doc.add_paragraph()
            run = paragraph.add_run(f"{i:04d}: {line.rstrip()}")
            font = run.font
            font.name = 'Courier New'
            font.size = Pt(9)
    else:
        doc.add_paragraph(f"[AVISO] Arquivo não encontrado: {path}", style='Intense Quote')

# Salvar o documento
output_path = "relatorios/appendice_scripts_codigo_fonte.docx"
os.makedirs(os.path.dirname(output_path), exist_ok=True)
doc.save(output_path)
print(f"Documento salvo com sucesso em: {output_path}")