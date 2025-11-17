# gerar_relatorio_pdf.py
# Gera relatório analítico final em PDF e DOCX com base nos dados processados e gráficos gerados pelo projeto.

import os
import pandas as pd
import matplotlib.pyplot as plt
from matplotlib.backends.backend_pdf import PdfPages
from docx import Document

def gerar_relatorio():
    df = pd.read_csv('dados/processado/dados_ingresso_evasao_conclusao.csv', sep=';')
    
    with PdfPages('relatorios/relatorio_analitico_final.pdf') as pdf:
        # Página 1: Estrutura do Relatório
        plt.figure(figsize=(8.5, 11))
        plt.axis('off')
        estrutura = [
            "1. INTRODUÇÃO",
            "2. OBJETIVOS",
            "  2.1 Objetivos gerais",
            "  2.2 Objetivos específicos",
            "    2.2.1 Cronograma do Projeto",
            "3. JUSTIFICATIVA",
            "4. REVISÃO BIBLIOGRÁFICA (SÍNTESE)",
            "5. METODOLOGIA",
            "  5.1 Definição dos objetivos",
            "  5.2 Seleção das bases de dados",
            "     • Microdados do Censo da Educação Superior – INEP: são arquivos CSV volumosos, divididos por ano.",
            "  5.3 Pré-processamento e transformação",
            "  5.4 Mineração dos dados, EDA, Análises, MVP",
            "  5.5 Interpretação e avaliação dos resultados",
            "6. RESULTADOS E DISCUSSÕES"
        ]
        for i, linha in enumerate(estrutura):
            plt.text(0.05, 1 - 0.05 * (i + 1), linha, fontsize=12)
        pdf.savefig()
        plt.close()

        # Página 2: Cronograma do Projeto
        cronograma_texto = [
            "CRONOGRAMA DE ATIVIDADES (2025)",
            "Definição do tema: Finalizado (13/03)",
            "Coleta de dados INEP: Até 13/03",
            "Pré-processamento e EDA: Até 13/03",
            "Validação cruzada e gráficos: Até 16/03",
            "Refinamento de análises: Até 22/03",
            "Validação com orientadora: 24/03",
            "Modelagem preditiva: 25/03 a 15/04",
            "Validação modelo: 10/04 a 16/04",
            "Otimização do modelo: 17/04 a 15/05",
            "Segunda entrega: Até 06/05",
            "Vídeo final: 06/05 a 20/05",
            "Entrega final: Até 25/05"
        ]
        plt.figure(figsize=(8.5, 11))
        plt.axis('off')
        for i, linha in enumerate(cronograma_texto):
            plt.text(0.05, 1 - 0.06 * (i + 1), linha, fontsize=12)
        pdf.savefig()
        plt.close()

        # Página 3: Introdução e Justificativa
        plt.figure(figsize=(8.5, 11))
        plt.axis('off')
        texto_intro = [
            "1. INTRODUÇÃO",
            "Este projeto tem como foco a análise de dados da educação superior brasileira,",
            "com ênfase na evasão e conclusão de cursos em instituições públicas e privadas.",
            "",
            "3. JUSTIFICATIVA",
            "A evasão acadêmica impacta diretamente o desenvolvimento educacional e a alocação",
            "de recursos públicos. Compreender padrões e prever tendências é fundamental para",
            "a tomada de decisões mais eficazes por parte de gestores e formuladores de políticas públicas."
        ]
        for i, linha in enumerate(texto_intro):
            plt.text(0.05, 1 - 0.06 * (i + 1), linha, fontsize=12, weight='bold' if i in [0, 4] else 'normal')
        pdf.savefig()
        plt.close()

        # Página 4: Metodologia
        plt.figure(figsize=(8.5, 11))
        plt.axis('off')
        metodologia = [
            "5. METODOLOGIA",
            "",
            "5.1 Definição dos objetivos",
            "- Estimar taxas de ingresso, evasão e conclusão de cursos superiores no Brasil.",
            "- Criar visualizações analíticas e um modelo preditivo para auxiliar diagnósticos.",
            "",
            "5.2 Seleção das bases de dados",
            "- Utilização dos Microdados do Censo da Educação Superior – INEP (1995–2023).",
            "- Arquivos CSV com milhões de registros, organizados anualmente.",
            "",
            "5.3 Pré-processamento e transformação",
            "- Leitura com pandas; tratamento de separadores, codificações e colunas inconsistentes.",
            "- Normalização com validações estatísticas e remoção de outliers.",
            "- Scripts utilizados: `pre_processamento.py`, `tratar_dados.py`.",
            "",
            "5.4 Mineração dos dados, EDA e MVP",
            "- Análises com `analises.py`, incluindo correlações e geração de gráficos (Seaborn, Matplotlib).",
            "- Scripts de visualização: `gerar_graficos.py` e `gerar_relatorio_pdf.py`.",
            "",
            "5.5 Modelagem e avaliação",
            "- Modelo preditivo de regressão linear implementado em `treinamento_modelo.py`.",
            "- Métricas avaliadas: erro quadrático médio e matriz de confusão.",
            "",
            "5.6 Tecnologias e ferramentas utilizadas",
            "- Linguagem: Python 3.11",
            "- Bibliotecas: pandas, numpy, seaborn, matplotlib, scikit-learn",
            "- Modelagem: Regressão Linear com validação cruzada (holdout)",
            "- Scripts de apoio: `pre_processamento.py`, `tratar_dados.py`, `analises.py`,",
            "  `treinamento_modelo.py`, `gerar_graficos.py`, `gerar_relatorio_pdf.py`",
            "- Versionamento: Git e GitHub (repositório com branches e commits por etapa)",
            "- Armazenamento e colaboração: Microsoft Teams (arquivos, reuniões e validações)",
            "- Ambiente de execução: macOS com VS Code, ambiente virtual (.venv)"
        ]
        for i, linha in enumerate(metodologia):
            plt.text(0.05, 1 - 0.045 * (i + 1), linha, fontsize=11, weight='bold' if linha.strip().endswith(":") or "5." in linha else 'normal')
        pdf.savefig()
        plt.close()

        # Outras páginas: inserir gráficos
        for nome in [
            "distribuicao_taxa_evasao", "grafico_taxa_conclusao",
            "grafico_taxa_ingresso", "grafico_ingressantes_vs_concluintes",
            "grafico_taxa_evasao_modalidade", "grafico_mapa_calor_correlacoes"
        ]:
            try:
                img_path = f"relatorios/figuras/{nome}.png"
                plt.figure(figsize=(10, 6))
                img = plt.imread(img_path)
                plt.imshow(img)
                plt.axis('off')
                plt.title(nome.replace("_", " ").title(), fontsize=14)
                pdf.savefig()
                plt.close()
            except FileNotFoundError:
                print(f"[AVISO] Imagem não encontrada: {img_path}")

        # Página final: Conclusão Analítica
        conclusoes = [
            "CONCLUSÃO GERAL",
            "- Taxa de evasão: Alta na maioria dos cursos, picos extremos em 100%",
            "- Taxa de conclusão: Muito baixa em diversos cursos",
            "- Relação ingressantes/concluintes: Correlação positiva, mas dispersa",
            "- Modalidade de ensino: EAD com risco crítico de evasão",
            "- Taxa de ingresso: Variação acentuada na ocupação de vagas",
            "- Modelagem: Regressão linear binária com 100% acurácia — revisar"
        ]
        plt.figure(figsize=(8.5, 11))
        plt.axis('off')
        for i, item in enumerate(conclusoes):
            plt.text(0.05, 1 - 0.08 * (i + 1), item, fontsize=12 if i > 0 else 14, weight='bold' if i == 0 else 'normal')
        pdf.savefig()
        plt.close()
        
        print("[OK] PDF gerado com sucesso.")

    # Geração do DOCX
    doc = Document()
    doc.add_heading('Relatório Analítico Final', 0)

    doc.add_heading('1. INTRODUÇÃO', level=1)
    doc.add_paragraph(
        "Este projeto tem como foco a análise de dados da educação superior brasileira, "
        "com ênfase na evasão e conclusão de cursos em instituições públicas e privadas."
    )

    doc.add_heading('2. OBJETIVOS', level=1)
    doc.add_heading('2.1 Objetivos gerais', level=2)
    doc.add_paragraph("- Estimar taxas de ingresso, evasão e conclusão de cursos superiores no Brasil.\n"
                      "- Criar visualizações analíticas e um modelo preditivo para auxiliar diagnósticos.")

    doc.add_heading('2.2 Objetivos específicos', level=2)
    doc.add_heading('2.2.1 Cronograma do Projeto', level=3)
    for item in cronograma_texto[1:]:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_heading('3. JUSTIFICATIVA', level=1)
    doc.add_paragraph(
        "A evasão acadêmica impacta diretamente o desenvolvimento educacional e a alocação de recursos públicos. "
        "Compreender padrões e prever tendências é fundamental para a tomada de decisões mais eficazes por parte de gestores e formuladores de políticas públicas."
    )

    doc.add_heading('4. REVISÃO BIBLIOGRÁFICA (SÍNTESE)', level=1)
    doc.add_paragraph("Conteúdo a ser incluído manualmente conforme revisão dos autores.")

    doc.add_heading('5. METODOLOGIA', level=1)
    for linha in metodologia:
        if linha.strip().endswith(":"):
            doc.add_heading(linha.strip(), level=2)
        elif linha.strip().startswith("-"):
            doc.add_paragraph(linha.strip(), style='List Bullet')
        elif linha.strip():
            doc.add_paragraph(linha.strip())

    doc.add_heading('6. RESULTADOS E DISCUSSÕES', level=1)
    for item in conclusoes[1:]:
        doc.add_paragraph(item, style='List Bullet')

    doc.save('relatorios/relatorio_analitico_final.docx')
    print("[OK] DOCX gerado com sucesso.")

if __name__ == '__main__':
    gerar_relatorio()